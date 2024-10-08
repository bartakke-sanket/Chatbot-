import streamlit as st
from PyPDF2 import PdfReader
from docx import Document
import pdfplumber
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_google_genai import GoogleGenerativeAIEmbeddings
import google.generativeai as genai
from langchain.vectorstores import FAISS
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.chains.question_answering import load_qa_chain
from langchain.prompts import PromptTemplate
from dotenv import load_dotenv
import os
from PIL import Image
import io
import traceback
# from exceptions import PendingDeprecationWarning


# Load environment variables
load_dotenv()

# Configure Google Generative AI API Key
api_key = os.getenv("GOOGLE_API_KEY")
if api_key:
    genai.configure(api_key=api_key)
else:
    st.error("GOOGLE_API_KEY not found in environment variables.")

# Document extraction functions
def extract_from_pdf(pdf_docs):
    text, tables, images = '', [], []
    for pdf in pdf_docs:
        with pdfplumber.open(pdf) as pdf_reader:
            for page in pdf_reader.pages:
                text += page.extract_text() or ''
                tables.extend(page.extract_tables() or [])
                images.extend(page.images)
    return text, tables, images

def extract_from_docx(docx_docs):
    text, tables, images = '', [], []
    for docx in docx_docs:
        doc = Document(docx)
        for para in doc.paragraphs:
            text += para.text + '\n'
        for table in doc.tables:
            tables.append([[cell.text for cell in row.cells] for row in table.rows])
    return text, tables, images

def get_text_chunks(text):
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=10000, chunk_overlap=2000)
    return text_splitter.split_text(text)

def get_vector_store(text_chunks):
    try:
        embeddings = GoogleGenerativeAIEmbeddings(model='models/embedding-001')
        vector_store = FAISS.from_texts(text_chunks, embedding=embeddings)
        vector_store.save_local('faiss_index')
        return vector_store
    except Exception as e:
        st.error(f"Error creating vector store: {str(e)}")
        st.error(traceback.format_exc())
        return None

def get_conversational_chain():
    prompt_template = """
    You are chatting with a document that may contain text, tables, and images.
    Answer questions based on the provided context. For tables, respond in tabular format.
    If an image is asked about, provide the image.
    If the answer is not in the context, say 'Answer not available in the provided context.'
    Do not provide incorrect answers.
    
    Context: {context}
    
    Question: {question}
    
    Answer:
    """
    model = ChatGoogleGenerativeAI(model="gemini-pro", temperature=0.3)
    prompt = PromptTemplate(template=prompt_template, input_variables=["context", "question"])
    return load_qa_chain(model, chain_type="stuff", prompt=prompt)

def user_input(user_question, vector_store):
    try:
        docs = vector_store.similarity_search(user_question, k=4)
        chain = get_conversational_chain()
        response = chain({"input_documents": docs, "question": user_question}, return_only_outputs=True)
        return response["output_text"]
    except Exception as e:
        st.error(f"An error occurred in user_input: {str(e)}")
        st.error(traceback.format_exc())
        return "An error occurred while processing your question. Please try again."

# Improved Chat UI
def main():
    st.set_page_config(page_title="Chat with Multiple Documents", page_icon=":books:", layout="wide")

    st.title("Chat with Your Documents :books:")

    if "conversation" not in st.session_state:
        st.session_state.conversation = []
    if "vector_store" not in st.session_state:
        st.session_state.vector_store = None

    with st.sidebar:
        st.subheader("Upload Your Documents (PDF or DOCX)")
        docs = st.file_uploader("Upload PDFs or Word Documents", type=["pdf", "docx"], accept_multiple_files=True)

        if st.button("Process"):
            if docs:
                with st.spinner("Processing your documents..."):
                    try:
                        text, tables, images = "", [], []
                        for doc in docs:
                            if doc.type == "application/pdf":
                                doc_text, doc_tables, doc_images = extract_from_pdf([doc])
                            elif doc.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                                doc_text, doc_tables, doc_images = extract_from_docx([doc])
                            text += doc_text
                            tables += doc_tables
                            images += doc_images

                        text_chunks = get_text_chunks(text)
                        st.session_state.vector_store = get_vector_store(text_chunks)

                        display_tables(tables)
                        display_images(images)

                        if st.session_state.vector_store:
                            st.success("Documents processed successfully. You can now ask questions!")
                        else:
                            st.error("Failed to process documents. Please try again.")
                    except Exception as e:
                        st.error(f"Error processing documents: {str(e)}")
                        st.error(traceback.format_exc())
            else:
                st.error("Please upload PDF or DOCX files before processing.")

    st.subheader("Chat with Your Documents")

    # Define custom CSS for improved UI with new colors
    st.markdown("""
        <style>
        .user-message {
            background-color: #00796B; /* Dark Teal */
            color: white;
            border-radius: 20px;
            padding: 10px;
            margin: 5px;
            text-align: right;
            width: fit-content;
            max-width: 70%;
            float: right;
            clear: both;
        }
        .assistant-message {
            background-color: #F57C00; /* Orange */
            color: white;
            border-radius: 20px;
            padding: 10px;
            margin: 5px;
            text-align: left;
            width: fit-content;
            max-width: 70%;
            float: left;
            clear: both;
        }
        .chat-container {
            max-height: 400px;
            overflow-y: auto;
            padding: 20px;
        }
        .chat-input {
            position: fixed;
            bottom: 0;
            width: 100%;
            padding: 10px;
            background-color: #FFFFFF;
        }
        .icon-bot {
            width: 40px;
            height: 40px;
            margin-right: 10px;
            margin-top: 5px;
        }
        .icon-user {
            width: 40px;
            height: 40px;
            margin-left: 10px;
            margin-top: 5px;
        }
        </style>
    """, unsafe_allow_html=True)

    # Display conversation history with icons and styled backgrounds
    st.markdown('<div class="chat-container">', unsafe_allow_html=True)
    for message in st.session_state.conversation:
        if message["role"] == "assistant":
            st.markdown(f"""
                <div class="assistant-message">
                    <img src="https://cdn-icons-png.flaticon.com/512/4712/4712021.png" class="icon-bot">
                    {message["content"]}
                </div>
            """, unsafe_allow_html=True)
        else:
            st.markdown(f"""
                <div class="user-message">
                    {message["content"]}
                    <img src="https://cdn-icons-png.flaticon.com/512/1177/1177568.png" class="icon-user">
                </div>
            """, unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)

    # Chat input box
    if st.session_state.vector_store:
        user_question = st.chat_input("Ask a question about your documents...")
        if user_question:
            st.session_state.conversation.append({"role": "human", "content": user_question})

            with st.spinner("Thinking..."):
                response = user_input(user_question, st.session_state.vector_store)

            st.session_state.conversation.append({"role": "assistant", "content": response})
            st.rerun()
    else:
        st.warning("Please upload and process documents before asking questions.")

# Functions for displaying tables and images
def display_tables(tables):
    if tables:
        st.write("### Extracted Tables:")
        for table in tables:
            st.table(table)
    else:
        st.write("No tables found.")

def display_images(images):
    if images:
        st.write("### Extracted Images:")
        for img in images:
            if isinstance(img, dict):
                st.write(f"Image properties: {img}")
            else:
                try:
                    img_pil = Image.open(io.BytesIO(img))
                    st.image(img_pil, caption="Extracted Image")
                except Exception as e:
                    st.write(f"Error displaying image: {e}")
    else:
        st.write("No images found.")

if __name__ == "__main__":
    main()
